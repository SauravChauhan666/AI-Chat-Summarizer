# app.py - Enhanced Professional Version
import os
import re
import json
import fitz            # PyMuPDF
import docx
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from openai import OpenAI
import logging
from typing import List, Dict, Optional

# -------------------- Config & Setup --------------------
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Environment validation
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    logger.error("OPENAI_API_KEY not found in environment")
    raise RuntimeError("Please set OPENAI_API_KEY in your environment or .env file.")

client = OpenAI(api_key=OPENAI_API_KEY)

# Directory setup
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
CHAT_HISTORY_FOLDER = os.path.join(BASE_DIR, "chat_history")
SAVED_CONV_FOLDER = os.path.join(BASE_DIR, "saved_conversations")
TEMPLATES_FOLDER = os.path.join(BASE_DIR, "templates")

# Create directories if they don't exist
for directory in (UPLOAD_FOLDER, CHAT_HISTORY_FOLDER, SAVED_CONV_FOLDER, TEMPLATES_FOLDER):
    os.makedirs(directory, exist_ok=True)

# File handling configuration
ALLOWED_EXT = {".pdf", ".docx", ".doc", ".txt"}
MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB

# Flask app configuration
app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "your-secret-key-change-this-in-production")
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = MAX_FILE_SIZE

# -------------------- Enhanced Helper Functions --------------------

def allowed_file(filename: str) -> bool:
    """Check if the file extension is allowed."""
    return os.path.splitext(filename.lower())[1] in ALLOWED_EXT

def validate_file_size(file_path: str) -> bool:
    """Validate file size is within limits."""
    try:
        return os.path.getsize(file_path) <= MAX_FILE_SIZE
    except OSError:
        return False

def read_file_content(path: str) -> str:
    """
    Enhanced file reading with better error handling and logging.
    Supports PDF, DOCX/DOC, and TXT files.
    """
    if not os.path.exists(path):
        logger.error(f"File not found: {path}")
        return "[Error: File not found]"
    
    if not validate_file_size(path):
        logger.error(f"File too large: {path}")
        return "[Error: File too large]"
    
    ext = os.path.splitext(path)[1].lower()
    text = ""
    
    try:
        if ext == ".pdf":
            logger.info(f"Reading PDF file: {path}")
            doc = fitz.open(path)
            pages = []
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text()
                    if page_text.strip():  # Only add non-empty pages
                        pages.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error reading page {page_num + 1}: {str(e)}")
                    continue
            text = "\n\n".join(pages)
            doc.close()
            
        elif ext in (".docx", ".doc"):
            logger.info(f"Reading DOCX/DOC file: {path}")
            doc = docx.Document(path)
            paragraphs = []
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:  # Skip empty paragraphs
                    paragraphs.append(para_text)
            text = "\n\n".join(paragraphs)
            
            # Also extract text from tables if present
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text += f"\n\n--- Table Data ---\n" + "\n".join(table_text)
                    
        elif ext == ".txt":
            logger.info(f"Reading TXT file: {path}")
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(path, "r", encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, read as binary and replace errors
                with open(path, "rb") as f:
                    text = f.read().decode('utf-8', errors='replace')
                    
    except Exception as e:
        logger.error(f"Error reading file {path}: {str(e)}")
        text = f"[Error reading file: {str(e)}]"
    
    # Clean and validate the extracted text
    if text:
        # Remove excessive whitespace but preserve structure
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 consecutive newlines
        text = text.strip()
        logger.info(f"Successfully extracted {len(text)} characters from {path}")
    else:
        logger.warning(f"No text extracted from {path}")
        text = "[No readable content found in file]"
    
    return text

def save_chat(chat_id: str, conversation_list: List[Dict]) -> bool:
    """Save chat with error handling and backup."""
    try:
        path = os.path.join(CHAT_HISTORY_FOLDER, f"{chat_id}.json")
        
        # Create backup if file exists
        if os.path.exists(path):
            backup_path = path + ".backup"
            os.rename(path, backup_path)
        
        with open(path, "w", encoding="utf-8") as f:
            json.dump(conversation_list, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Chat saved successfully: {chat_id}")
        return True
        
    except Exception as e:
        logger.error(f"Error saving chat {chat_id}: {str(e)}")
        return False

def load_chat(chat_id: str) -> List[Dict]:
    """Load chat with error handling."""
    try:
        path = os.path.join(CHAT_HISTORY_FOLDER, f"{chat_id}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            logger.info(f"Chat loaded successfully: {chat_id}")
            return data
    except Exception as e:
        logger.error(f"Error loading chat {chat_id}: {str(e)}")
    
    return []

def build_prompt(context_msgs: List[Dict], new_message: str, file_content: str) -> str:
    """
    Enhanced prompt builder with better context management and formatting.
    """
    # Create context from recent messages (limit to prevent token overflow)
    context_text = ""
    if context_msgs:
        # Get last 10 messages for context
        recent_msgs = context_msgs[-10:] if len(context_msgs) > 10 else context_msgs
        context_lines = []
        
        for i, msg in enumerate(recent_msgs):
            if isinstance(msg, dict):
                message_text = msg.get('message', '').strip()
                timestamp = msg.get('timestamp', '')
                if message_text:
                    # Format timestamp for readability
                    try:
                        if timestamp:
                            dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                            time_str = dt.strftime("%Y-%m-%d %H:%M")
                            context_lines.append(f"[{time_str}] {message_text[:200]}...")
                        else:
                            context_lines.append(f"[Message {i+1}] {message_text[:200]}...")
                    except:
                        context_lines.append(f"[Message {i+1}] {message_text[:200]}...")
        
        context_text = "\n".join(context_lines) if context_lines else "No prior conversation."
    else:
        context_text = "No prior conversation."

    # Enhanced prompt with better structure and instructions
    prompt = f"""You are an expert business meeting summarizer and AI assistant specializing in creating professional, investor-ready reports. Your task is to analyze meeting conversations and documents to produce structured, actionable summaries.

CRITICAL FORMATTING REQUIREMENTS:
- Use EXACTLY these section headers with colons: "Summary:", "Action Points:", "Dates & Deadlines:", "Document Summary:", "Detailed Summary:", "Confidence:"
- Each section must be clearly separated
- Use bullet points (•) for all lists
- Be concise but comprehensive
- Use professional business language
- Include specific names, numbers, and dates when mentioned

CONVERSATION CONTEXT (Recent Messages):
{context_text}

CURRENT MESSAGE TO ANALYZE:
{new_message}

ATTACHED DOCUMENT CONTENT:
{file_content if file_content and not file_content.startswith('[Error') else "No document attached or document unreadable."}

REQUIRED OUTPUT STRUCTURE:

Summary:
• Create 2-4 concise bullet points capturing key decisions, updates, and developments
• Focus on outcomes and high-level takeaways
• Use action-oriented language

Action Points:
• List specific tasks with format: "Task description — Assigned to [person/team] — Due [date]"
• If assignee or deadline not specified, omit those parts
• Prioritize actionable items
• Use bullet points (•) for each item

Dates & Deadlines:
• List all specific dates, deadlines, and time-sensitive items mentioned
• Format: "Event/Task — Date/Deadline"
• If no dates mentioned, state: "• No specific dates mentioned"

Document Summary:
• If document provided, summarize key findings, recommendations, or data in 2-5 bullets
• Focus on actionable insights and important information
• If no document: "• No documents uploaded"

Detailed Summary:
• Provide comprehensive breakdown of the discussion
• Include context, rationale, alternatives considered, and unresolved issues
• Use bullet points and sub-bullets for organization
• Capture nuances and important details that support decision-making

Confidence:
Provide a single line assessment like: "Confidence: [High/Medium/Low] — [brief reason]"

Remember: This summary may be shared with executives and investors. Ensure clarity, accuracy, and professionalism throughout."""

    return prompt

def call_openai_chat(prompt: str) -> str:
    """
    Enhanced OpenAI API call with retry logic and better error handling.
    """
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            logger.info(f"Calling OpenAI API (attempt {attempt + 1}/{max_retries})")
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",  # You can change to gpt-4 for better quality
                messages=[
                    {
                        "role": "system", 
                        "content": "You are a professional business meeting summarizer. Always follow the exact formatting requirements provided in the user prompt."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,  # Lower for more consistent formatting
                max_tokens=1500,  # Increased for detailed summaries
                top_p=0.95
            )
            
            content = response.choices[0].message.content
            logger.info(f"OpenAI API call successful, response length: {len(content)}")
            return content
            
        except Exception as e:
            logger.error(f"OpenAI API error (attempt {attempt + 1}): {str(e)}")
            if attempt == max_retries - 1:
                return f"Error generating summary: {str(e)}\n\nPlease try again or contact support if the issue persists."
    
    return "Failed to generate summary after multiple attempts."

def parse_sections(model_text: str) -> Dict[str, str]:
    """
    Enhanced section parser with better regex and fallback handling.
    """
    sections = {
        "Summary": "",
        "Action Points": "",
        "Dates & Deadlines": "",
        "Document Summary": "",
        "Detailed Summary": "",
        "Confidence": ""
    }

    if not model_text.strip():
        return sections

    text = model_text.strip()
    
    # Improved regex pattern for section headers
    pattern = re.compile(
        r"^(Summary|Action Points|Dates\s*&?\s*Deadlines|Document Summary|Detailed Summary|Confidence)\s*:\s*",
        flags=re.MULTILINE | re.IGNORECASE
    )
    
    # Split text by section headers
    parts = pattern.split(text)
    
    if len(parts) > 1:
        # Process matched sections
        current_section = None
        for i in range(1, len(parts), 2):  # Skip first part (before any header)
            if i < len(parts):
                header = parts[i].strip()
                content = parts[i + 1].strip() if i + 1 < len(parts) else ""
                
                # Normalize header names
                if header.lower().startswith('summary'):
                    sections["Summary"] = content
                elif header.lower().startswith('action'):
                    sections["Action Points"] = format_action_points(content)
                elif 'deadline' in header.lower() or 'date' in header.lower():
                    sections["Dates & Deadlines"] = content
                elif header.lower().startswith('document'):
                    sections["Document Summary"] = content
                elif header.lower().startswith('detailed'):
                    sections["Detailed Summary"] = content
                elif header.lower().startswith('confidence'):
                    sections["Confidence"] = content
    else:
        # Fallback: put entire text in Detailed Summary
        logger.warning("Could not parse sections, using fallback")
        sections["Detailed Summary"] = text

    return sections

def format_action_points(content: str) -> str:
    """Format action points with consistent bullet formatting."""
    if not content.strip():
        return content
    
    lines = []
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Ensure bullet point
        if not line.startswith(('•', '-', '*')):
            line = '• ' + line
        elif line.startswith(('-', '*')):
            line = '• ' + line[1:].strip()
        
        lines.append(line)
    
    return '\n'.join(lines)

def cleanup_old_files(days_old: int = 30):
    """Clean up old uploaded files to save disk space."""
    try:
        cutoff_time = datetime.now().timestamp() - (days_old * 24 * 3600)
        
        for folder in [UPLOAD_FOLDER]:
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                if os.path.isfile(file_path) and os.path.getctime(file_path) < cutoff_time:
                    os.remove(file_path)
                    logger.info(f"Cleaned up old file: {filename}")
                    
    except Exception as e:
        logger.error(f"Error during cleanup: {str(e)}")

# -------------------- Enhanced Routes --------------------

@app.route("/", methods=["GET", "POST"])
def home():
    """Enhanced home route with better error handling and user experience."""
    try:
        # Get existing chats sorted by modification time
        chat_files = []
        if os.path.exists(CHAT_HISTORY_FOLDER):
            for filename in os.listdir(CHAT_HISTORY_FOLDER):
                if filename.endswith(".json"):
                    file_path = os.path.join(CHAT_HISTORY_FOLDER, filename)
                    try:
                        mtime = os.path.getmtime(file_path)
                        chat_id = filename[:-5]  # Remove .json extension
                        chat_files.append((chat_id, mtime))
                    except OSError:
                        continue
        
        # Sort by modification time (newest first)
        chat_files.sort(key=lambda x: x[1], reverse=True)
        chat_files = [chat_id for chat_id, _ in chat_files]
        
        if request.method == "POST":
            return handle_summary_request(chat_files)
        
        # Handle GET request
        continue_chat_id = request.args.get('continue_chat_id')
        past_messages = []
        
        if continue_chat_id:
            past_messages = load_chat(continue_chat_id)
        
        return render_template(
            "home.html", 
            chat_files=chat_files[:20],  # Limit to recent 20 chats
            continue_chat_id=continue_chat_id,
            past_messages=past_messages[-10:] if past_messages else []  # Show last 10 messages
        )
        
    except Exception as e:
        logger.error(f"Error in home route: {str(e)}")
        flash("An error occurred. Please try again.", "danger")
        return render_template("home.html", chat_files=[])

def handle_summary_request(chat_files: List[str]):
    """Handle POST request for summary generation."""
    try:
        # Extract form data
        chat_option = request.form.get("chat_option", "new")
        chat_id = request.form.get("chat_id", "").strip()
        message = request.form.get("message", "").strip()
        file = request.files.get("file")

        # Validate inputs
        if not message:
            flash("Please enter meeting notes or conversation content.", "danger")
            return redirect(url_for("home"))

        # Generate chat ID if not provided
        if not chat_id:
            chat_id = datetime.now().strftime("%Y%m%d%H%M%S")

        # Load existing conversation
        conversation = load_chat(chat_id) if chat_option == "continue" else []
        
        # Handle file upload
        file_content = ""
        if file and file.filename:
            if not allowed_file(file.filename):
                flash("File type not allowed. Supported: PDF, DOCX, DOC, TXT", "danger")
                return redirect(url_for("home"))
            
            # Save file with timestamp
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            saved_filename = f"{timestamp}_{filename}"
            saved_path = os.path.join(app.config["UPLOAD_FOLDER"], saved_filename)
            
            try:
                file.save(saved_path)
                file_content = read_file_content(saved_path)
                logger.info(f"File uploaded and processed: {saved_filename}")
            except Exception as e:
                logger.error(f"File upload error: {str(e)}")
                flash("Error uploading file. Please try again.", "danger")
                return redirect(url_for("home"))

        # Build prompt and generate summary
        try:
            prompt = build_prompt(conversation, message, file_content)
            model_output = call_openai_chat(prompt)
            sections = parse_sections(model_output)
            
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            flash("Error generating summary. Please try again.", "danger")
            return redirect(url_for("home"))

        # Save to conversation history
        conversation_entry = {
            "timestamp": datetime.utcnow().isoformat() + "Z",
            "message": message,
            "file_attached": bool(file and file.filename),
            "file_name": file.filename if file and file.filename else None,
            "model_output": model_output,
            "sections": sections
        }
        
        conversation.append(conversation_entry)
        
        if not save_chat(chat_id, conversation):
            flash("Warning: Could not save conversation history.", "warning")

        # Prepare data for template
        conversation_data = json.dumps(conversation, ensure_ascii=False)

        return render_template(
            "result.html",
            chat_id=chat_id,
            sections=sections,
            raw_text=model_output,
            conversation_data=conversation_data,
            chat_files=chat_files
        )

    except Exception as e:
        logger.error(f"Error handling summary request: {str(e)}")
        flash("An unexpected error occurred. Please try again.", "danger")
        return redirect(url_for("home"))

@app.route("/continue/<chat_id>", methods=["GET"])
def continue_chat(chat_id: str):
    """Continue an existing chat conversation."""
    try:
        conversation = load_chat(chat_id)
        chat_files = []
        
        # Get chat files list
        try:
            chat_files = sorted(
                [f[:-5] for f in os.listdir(CHAT_HISTORY_FOLDER) if f.endswith(".json")], 
                reverse=True
            )
        except OSError:
            pass
        
        return render_template(
            "home.html", 
            chat_files=chat_files,
            continue_chat_id=chat_id, 
            past_messages=conversation[-10:] if conversation else []
        )
        
    except Exception as e:
        logger.error(f"Error continuing chat {chat_id}: {str(e)}")
        flash("Error loading chat history.", "danger")
        return redirect(url_for("home"))

@app.route("/save_conversation", methods=["POST"])
def save_conversation():
    """Enhanced conversation saving with validation."""
    try:
        conversation_name = request.form.get("conversation_name", "").strip()
        conversation_data = request.form.get("conversation_data", "")
        
        if not conversation_name:
            flash("Please provide a name to save the conversation.", "danger")
            return redirect(request.referrer or url_for("home"))

        # Sanitize filename
        safe_name = secure_filename(conversation_name)
        if not safe_name:
            flash("Invalid file name. Please use alphanumeric characters.", "danger")
            return redirect(request.referrer or url_for("home"))

        # Create unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        save_path = os.path.join(SAVED_CONV_FOLDER, f"{safe_name}_{timestamp}.json")
        
        # Prepare data for saving
        save_data = {
            "name": conversation_name,
            "saved_at": datetime.utcnow().isoformat() + "Z",
            "conversation_data": json.loads(conversation_data) if conversation_data else []
        }

        with open(save_path, "w", encoding="utf-8") as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)

        logger.info(f"Conversation saved: {safe_name}")
        flash(f"Conversation saved successfully as '{conversation_name}'", "success")
        
    except json.JSONDecodeError:
        logger.error("Invalid conversation data format")
        flash("Error saving conversation: Invalid data format.", "danger")
    except Exception as e:
        logger.error(f"Error saving conversation: {str(e)}")
        flash("Error saving conversation. Please try again.", "danger")
    
    return redirect(request.referrer or url_for("home"))

@app.route("/api/chat/<chat_id>")
def api_get_chat(chat_id: str):
    """API endpoint to get chat data in JSON format."""
    try:
        conversation = load_chat(chat_id)
        return jsonify({
            "success": True,
            "chat_id": chat_id,
            "messages": conversation,
            "message_count": len(conversation)
        })
    except Exception as e:
        logger.error(f"Error retrieving chat {chat_id} via API: {str(e)}")
        return jsonify({
            "success": False,
            "error": "Chat not found or error occurred"
        }), 404

@app.route("/api/chats")
def api_list_chats():
    """API endpoint to list all available chats."""
    try:
        chat_files = []
        if os.path.exists(CHAT_HISTORY_FOLDER):
            for filename in os.listdir(CHAT_HISTORY_FOLDER):
                if filename.endswith(".json"):
                    chat_id = filename[:-5]
                    file_path = os.path.join(CHAT_HISTORY_FOLDER, filename)
                    try:
                        mtime = os.path.getmtime(file_path)
                        # Get first message preview
                        conversation = load_chat(chat_id)
                        preview = ""
                        if conversation:
                            first_msg = conversation[0].get('message', '')
                            preview = first_msg[:100] + "..." if len(first_msg) > 100 else first_msg
                        
                        chat_files.append({
                            "chat_id": chat_id,
                            "last_modified": datetime.fromtimestamp(mtime).isoformat(),
                            "message_count": len(conversation),
                            "preview": preview
                        })
                    except (OSError, json.JSONDecodeError):
                        continue
        
        # Sort by modification time
        chat_files.sort(key=lambda x: x['last_modified'], reverse=True)
        
        return jsonify({
            "success": True,
            "chats": chat_files,
            "total_chats": len(chat_files)
        })
        
    except Exception as e:
        logger.error(f"Error listing chats via API: {str(e)}")
        return jsonify({
            "success": False,
            "error": "Error retrieving chat list"
        }), 500

@app.route("/health")
def health_check():
    """Health check endpoint for monitoring."""
    try:
        # Basic health checks
        checks = {
            "upload_folder": os.path.exists(UPLOAD_FOLDER) and os.access(UPLOAD_FOLDER, os.W_OK),
            "chat_history_folder": os.path.exists(CHAT_HISTORY_FOLDER) and os.access(CHAT_HISTORY_FOLDER, os.W_OK),
            "saved_conv_folder": os.path.exists(SAVED_CONV_FOLDER) and os.access(SAVED_CONV_FOLDER, os.W_OK),
            "openai_api_key": bool(OPENAI_API_KEY),
            "templates_folder": os.path.exists(TEMPLATES_FOLDER)
        }
        
        all_healthy = all(checks.values())
        
        return jsonify({
            "status": "healthy" if all_healthy else "unhealthy",
            "checks": checks,
            "timestamp": datetime.utcnow().isoformat() + "Z"
        }), 200 if all_healthy else 503
        
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return jsonify({
            "status": "error",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat() + "Z"
        }), 500

@app.route("/delete_chat/<chat_id>", methods=["POST"])
def delete_chat(chat_id: str):
    """Delete a chat conversation."""
    try:
        chat_path = os.path.join(CHAT_HISTORY_FOLDER, f"{chat_id}.json")
        
        if os.path.exists(chat_path):
            os.remove(chat_path)
            logger.info(f"Chat deleted: {chat_id}")
            flash("Chat deleted successfully.", "success")
        else:
            flash("Chat not found.", "warning")
            
    except Exception as e:
        logger.error(f"Error deleting chat {chat_id}: {str(e)}")
        flash("Error deleting chat.", "danger")
    
    return redirect(url_for("home"))

@app.route("/export_chat/<chat_id>")
def export_chat(chat_id: str):
    """Export chat as JSON download."""
    try:
        conversation = load_chat(chat_id)
        if not conversation:
            flash("Chat not found.", "warning")
            return redirect(url_for("home"))
        
        from flask import make_response
        
        export_data = {
            "chat_id": chat_id,
            "exported_at": datetime.utcnow().isoformat() + "Z",
            "conversation": conversation
        }
        
        response = make_response(json.dumps(export_data, ensure_ascii=False, indent=2))
        response.headers['Content-Type'] = 'application/json'
        response.headers['Content-Disposition'] = f'attachment; filename=chat_{chat_id}.json'
        
        return response
        
    except Exception as e:
        logger.error(f"Error exporting chat {chat_id}: {str(e)}")
        flash("Error exporting chat.", "danger")
        return redirect(url_for("home"))

# -------------------- Error Handlers --------------------

@app.errorhandler(413)
def too_large(e):
    flash("File too large. Maximum file size is 25MB.", "danger")
    return redirect(url_for("home"))

@app.errorhandler(404)
def page_not_found(e):
    return render_template("error.html", 
                         error_code=404, 
                         error_message="Page not found"), 404

@app.errorhandler(500)
def internal_server_error(e):
    logger.error(f"Internal server error: {str(e)}")
    return render_template("error.html", 
                         error_code=500, 
                         error_message="Internal server error"), 500

# -------------------- Startup Tasks --------------------

def initialize_app():
    """Initialize application with startup tasks."""
    logger.info("Initializing Business Meeting Summarizer...")
    
    # Clean up old files on startup
    try:
        cleanup_old_files(30)  # Clean files older than 30 days
    except Exception as e:
        logger.warning(f"Cleanup failed: {str(e)}")
    
    # Test OpenAI connection
    try:
        test_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": "Test connection"}],
            max_tokens=10
        )
        logger.info("OpenAI API connection test successful")
    except Exception as e:
        logger.error(f"OpenAI API connection test failed: {str(e)}")
        
    logger.info("Application initialized successfully")

# -------------------- Main Application --------------------

if __name__ == "__main__":
    initialize_app()
    
    # Development vs Production settings
    debug_mode = os.getenv("FLASK_ENV") == "development"
    port = int(os.getenv("PORT", 5000))
    host = os.getenv("HOST", "0.0.0.0")
    
    if debug_mode:
        logger.info("Running in development mode")
        app.run(debug=True, host=host, port=port)
    else:
        logger.info("Running in production mode")
        # For production, use a proper WSGI server like gunicorn
        # gunicorn --bind 0.0.0.0:5000 --workers 4 app:app
        app.run(debug=False, host=host, port=port)